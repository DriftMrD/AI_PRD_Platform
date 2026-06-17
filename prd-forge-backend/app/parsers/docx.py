"""DOCX 解析器（python-docx 累加段落文本）。"""

from __future__ import annotations

import io

from docx import Document  # type: ignore[import-untyped]

from app.errors import AppError, FILE_PARSE_FAILED
from app.parsers._truncate import truncate_with_marker


class DocxParser:
    """`.docx` 解析：遍历 `paragraph.text`。"""

    def parse(self, content: bytes, filename: str) -> str:
        try:
            doc = Document(io.BytesIO(content))
        except Exception as exc:  # python-docx 抛 PackageNotFoundError 等
            raise AppError(
                FILE_PARSE_FAILED,
                f"DOCX 解析失败：{filename}：{exc}",
                retriable=False,
            ) from exc

        lines: list[str] = [p.text for p in doc.paragraphs]

        # 同时遍历表格（按行/按格累加，避免遗漏结构化内容）
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                lines.append(" | ".join(cells))

        text = "\n".join(lines).strip()
        return truncate_with_marker(text, filename)
